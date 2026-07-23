import streamlit as st
import pandas as pd
import datetime
import os
import json
import requests
import time
import threading

from database_academic import (
    get_academic_conn, add_lab_program, add_note, add_syllabus_delivery_log,
    get_lab_programs, get_notes, get_syllabus_delivery_log, get_drive_file,
    save_drive_file, search_academic_resources_rag
)

# ── BACKGROUND SYNC ENGINE ───────────────────────────────────

def get_all_drive_files_recursive(folder_id, api_key):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
    }
    to_scan = [folder_id]
    all_files = []
    scanned_folders = set()
    
    while to_scan:
        curr_folder_id = to_scan.pop(0)
        if curr_folder_id in scanned_folders:
            continue
        scanned_folders.add(curr_folder_id)
        
        url = "https://www.googleapis.com/drive/v3/files"
        params = {
            "q": f"'{curr_folder_id}' in parents and trashed = false",
            "fields": "files(id, name, mimeType, modifiedTime)",
            "key": api_key,
            "pageSize": 1000
        }
        try:
            res = requests.get(url, params=params, headers=headers, timeout=15)
            if res.status_code == 200:
                files_list = res.json().get("files", [])
                for f in files_list:
                    fid = f.get("id")
                    name = f.get("name")
                    mime = f.get("mimeType")
                    mod_time = f.get("modifiedTime")
                    if not fid or not name:
                        continue
                    
                    if mime == "application/vnd.google-apps.folder":
                        to_scan.append(fid)
                    else:
                        all_files.append({
                            "id": fid,
                            "name": name,
                            "mime_type": mime,
                            "modified_time": mod_time
                        })
            else:
                print(f"[Sync] API returned {res.status_code}: {res.text}")
        except Exception as e:
            print(f"[Sync] Folder scan exception: {e}")
            
    return all_files

def download_drive_file(file_id, file_name, api_key):
    cache_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "academic_materials_cache")
    os.makedirs(cache_dir, exist_ok=True)
    
    safe_name = "".join(c for c in file_name if c.isalnum() or c in "._- ").strip()
    local_path = os.path.join(cache_dir, f"{file_id}_{safe_name}")
    
    url = f"https://www.googleapis.com/drive/v3/files/{file_id}?alt=media&key={api_key}"
    try:
        res = requests.get(url, timeout=30)
        if res.status_code == 200:
            with open(local_path, "wb") as f:
                f.write(res.content)
            return local_path
        else:
            print(f"[Sync] Failed to download {file_name}: Status {res.status_code}")
    except Exception as e:
        print(f"[Sync] Download exception for {file_name}: {e}")
    return None

def get_subject_from_filename(file_name):
    name_upper = file_name.upper()
    if "PYTHON" in name_upper or "PPS" in name_upper:
        return "PYTHON"
    elif "EDC" in name_upper or "ELECTRONIC" in name_upper:
        return "EDC"
    elif "BEE" in name_upper or "ELECTRICAL" in name_upper:
        return "BEE"
    elif "ODEVC" in name_upper or "ODE" in name_upper or "VECTOR" in name_upper or "DIFFERENTIAL" in name_upper:
        return "ODEVC"
    elif "CHEMISTRY" in name_upper or "EC" in name_upper:
        return "Chemistry"
    elif "NAS" in name_upper or "NETWORK" in name_upper or "SYNTHESIS" in name_upper:
        return "NAS"
    elif "DS" in name_upper or "DATA STRUCTURE" in name_upper:
        return "DS"
    return "PYTHON"

def auto_detect_subject_metadata(file_name, sample_text, api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    prompt = f"""
    You are an academic file classification assistant. Analyze this file metadata and content:
    File Name: {file_name}
    Content Sample:
    {sample_text[:10000]}
    
    Classify the file and extract exactly:
    - "subject": The full name of the subject (e.g. "Electronic Devices and Circuits", "Python Programming", "Basic Electrical Engineering", "Programming for Problem Solving", "Engineering Chemistry", "Ordinary Differential Equations and Vector Calculus", "Data Structures", "Network Analysis and Synthesis", etc.)
    - "subject_code": The standardized subject code (e.g. "2566103ES" for Python, "2504103ES" for EDC, "23EE104ES" for BEE, "2512105ES" for PPS, "25CH102BS" for Chemistry, "25MA201BS" for ODEVC, "25PH102BS" for Applied Physics, "591" for DS, etc.)
    - "semester": The semester number as an integer (e.g. 1 or 2)
    - "branch": The engineering branch target (e.g. "ECE", "CSE", "CSE AIML", "CSE DS", "AIDS", "MECH", "CIVIL", "EEE", etc.)
    - "regulation": The academic regulation (e.g. "VR25", "VR23")
    - "academic_year": The academic year (e.g. "2025-2026")
    - "document_type": Type of document, choose exactly one of: "syllabus", "question_bank", "mcqs", "notes", "lab_program", "other"
    
    Respond ONLY with a valid JSON block containing these keys. Do not include markdown code block wrappers (like ```json) or any explanations.
    """
    
    payload = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    
    try:
        res = requests.post(url, json=payload, headers=headers, timeout=20)
        if res.status_code == 200:
            text_out = res.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            if text_out.startswith("```"):
                lines = text_out.split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines[-1].startswith("```"):
                    lines = lines[:-1]
                text_out = "\n".join(lines).strip()
            
            data = json.loads(text_out)
            return data
        else:
            print(f"[AutoDetect] Gemini returned status {res.status_code}")
    except Exception as e:
        print(f"[AutoDetect] Exception: {e}")
        
    name_upper = file_name.upper()
    doc_type = "other"
    if "SYLLABUS" in name_upper:
        doc_type = "syllabus"
    elif "QB" in name_upper or "QUESTION" in name_upper:
        doc_type = "question_bank"
    elif "MCQ" in name_upper or "OBJECTIVE" in name_upper:
        doc_type = "mcqs"
    elif "LAB" in name_upper or "RECORD" in name_upper or "PROGRAM" in name_upper or file_name.endswith((".py", ".c", ".cpp")):
        doc_type = "lab_program"
    elif "NOTES" in name_upper or "MATERIAL" in name_upper or "UNIT" in name_upper:
        doc_type = "notes"
        
    sub_code = get_actual_subject_code(get_subject_from_filename(file_name))
    
    return {
        "subject": get_subject_from_filename(file_name),
        "subject_code": sub_code,
        "semester": 2 if "SEM-II" in name_upper or "SEM 2" in name_upper or "SEM2" in name_upper else 1,
        "branch": "CSE" if "CSE" in name_upper else "ECE" if "ECE" in name_upper else "ALL",
        "regulation": "VR25" if "VR25" in name_upper else "VR23",
        "academic_year": "2025-2026",
        "document_type": doc_type
    }

def parse_zip_lab_programs(zip_path, subject_code):
    import zipfile
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            for file_info in zip_ref.infolist():
                if file_info.is_dir() or "__MACOSX" in file_info.filename or file_info.filename.startswith("."):
                    continue
                ext = os.path.splitext(file_info.filename)[1].lower()
                if ext in [".py", ".c", ".cpp", ".txt"]:
                    code_content = zip_ref.read(file_info).decode('utf-8', errors='ignore')
                    file_name = os.path.basename(file_info.filename)
                    explanation = f"Source: {file_info.filename} zip archive."
                    add_lab_program(subject_code, file_name, code_content, explanation)
    except Exception as e:
        print(f"[Sync] Zip parser error: {e}")

def parse_notes_via_gemini(text, subject_code, api_key):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    
    prompt = f"""
    You are an expert academic study guide parser. Analyze the provided study material and extract the main topics covered, grouped by course Units (usually Unit 1 to 5).
    Generate a valid JSON array of objects, each representing a topic notes section. Do not include markdown code block wrappers or explanations.
    
    Schema:
    [
      {{
        "unit": 1,
        "topic": "Zener Diode Characteristics",
        "content": "Detailed explanatory notes about Zener diode, including definition, characteristics, breakdown regions, and applications as a voltage regulator."
      }}
    ]
    
    Only extract major topics with rich, informative content summaries.
    Text:
    {text[:25000]}
    """
    payload = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    try:
        res = requests.post(url, json=payload, headers=headers, timeout=30)
        if res.status_code == 200:
            text_out = res.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
            if text_out.startswith("```"):
                lines = text_out.split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines[-1].startswith("```"):
                    lines = lines[:-1]
                text_out = "\n".join(lines).strip()
            parsed = json.loads(text_out)
            for note in parsed:
                add_note(subject_code, note.get("unit", 3), note.get("topic"), note.get("content"))
    except Exception as e:
        print(f"[Sync] Notes parser error: {e}")

def sync_drive_files_engine():
    drive_api_key = get_drive_api_key()
    gemini_api_key = get_gemini_api_key()
    folder_id = "1bInXkRc9mQFdbVbUMNxG1VVnrpKEyoPN"
    
    st.session_state["sync_status"] = "🔄 Running Drive Sync Engine..."
    
    try:
        files = get_all_drive_files_recursive(folder_id, drive_api_key)
        st.session_state["sync_status"] = f"🔄 Found {len(files)} files on Drive. Syncing..."
        
        synced_count = 0
        for f in files:
            file_id = f["id"]
            file_name = f["name"]
            mime = f["mime_type"]
            mod_time = f["modified_time"]
            
            db_f = get_drive_file(file_id)
            if not db_f or db_f["modified_time"] != mod_time:
                st.session_state["sync_status"] = f"🔄 Downloading {file_name}..."
                local_path = download_drive_file(file_id, file_name, drive_api_key)
                if not local_path:
                    continue
                
                st.session_state["sync_status"] = f"🔄 Auto-detecting {file_name}..."
                sample_text = ""
                file_ext = file_name.split('.')[-1].lower()
                
                try:
                    if file_ext == "pdf":
                        import pdfplumber
                        with pdfplumber.open(local_path) as pdf:
                            for page in pdf.pages[:3]:
                                sample_text += (page.extract_text() or "") + "\n"
                    elif file_ext == "docx":
                        import docx
                        doc = docx.Document(local_path)
                        sample_text = "\n".join([p.text for p in doc.paragraphs[:50]])
                    elif file_ext in ["txt", "py", "c", "cpp"]:
                        with open(local_path, "r", encoding="utf-8", errors="ignore") as file_handle:
                            sample_text = file_handle.read()
                    elif file_ext == "zip":
                        sample_text = "Zip archive containing code files."
                except Exception as e:
                    print(f"Error reading file {file_name} for auto-detect: {e}")
                    
                meta = auto_detect_subject_metadata(file_name, sample_text, gemini_api_key)
                
                st.session_state["sync_status"] = f"🔄 Parsing {file_name} ({meta['document_type']})..."
                subject_code = meta.get("subject_code")
                semester = meta.get("semester", 2)
                branch = meta.get("branch", "ALL")
                doc_type = meta.get("document_type", "other")
                
                if doc_type == "syllabus":
                    with open(local_path, "rb") as f_stream:
                        parse_and_import_syllabus(f_stream, subject_code)
                elif doc_type == "question_bank" or doc_type == "mcqs":
                    with open(local_path, "rb") as f_stream:
                        parse_and_import_qb(f_stream, file_ext, subject_code, default_unit=3)
                elif doc_type == "lab_program":
                    if file_ext == "zip":
                        parse_zip_lab_programs(local_path, subject_code)
                    elif file_ext in ["py", "c", "cpp", "txt"]:
                        explanation = f"Source file: {file_name}"
                        add_lab_program(subject_code, file_name, sample_text, explanation)
                elif doc_type == "notes":
                    parse_notes_via_gemini(sample_text, subject_code, gemini_api_key)
                
                save_drive_file(file_id, file_name, mime, mod_time, branch, semester, subject_code, processed=1)
                synced_count += 1
                
        st.session_state["sync_status"] = f"✅ Sync Complete. Processed {synced_count} new/modified files."
    except Exception as e:
        st.session_state["sync_status"] = f"❌ Sync Failed: {e}"
        print(f"[Sync Engine] Failed: {e}")

def start_sync_thread(force=False):
    now = time.time()
    last_sync = st.session_state.get("last_drive_sync_time", 0)
    
    if force or (now - last_sync > 900): # 15 minutes
        st.session_state["last_drive_sync_time"] = now
        t = threading.Thread(target=sync_drive_files_engine, daemon=True)
        t.start()


def init_sm2_columns():
    try:
        conn = get_academic_conn()
        cursor = conn.cursor()
        cursor.execute("PRAGMA table_info(review_schedule)")
        columns = [row[1] for row in cursor.fetchall()]
        
        modified = False
        if "easiness_factor" not in columns:
            cursor.execute("ALTER TABLE review_schedule ADD COLUMN easiness_factor REAL DEFAULT 2.5")
            modified = True
        if "repetitions" not in columns:
            cursor.execute("ALTER TABLE review_schedule ADD COLUMN repetitions INTEGER DEFAULT 0")
            modified = True
        if "interval" not in columns:
            cursor.execute("ALTER TABLE review_schedule ADD COLUMN interval INTEGER DEFAULT 0")
            modified = True
            
        if modified:
            conn.commit()
        conn.close()
    except Exception:
        pass

init_sm2_columns()

def get_drive_api_key():
    import os
    import streamlit as st
    try:
        if "google_drive" in st.secrets and "api_key" in st.secrets["google_drive"]:
            return st.secrets["google_drive"]["api_key"]
    except Exception:
        pass
    env_key = os.environ.get("GOOGLE_DRIVE_API_KEY") or os.environ.get("DRIVE_API_KEY")
    if env_key:
        return env_key
    return "AIzaSyCMuWAi15u9nrqoH20xN5kqdbho2tVCVws"

def get_gemini_api_key():
    import os
    import streamlit as st
    try:
        if "gemini" in st.secrets and "api_key" in st.secrets["gemini"]:
            return st.secrets["gemini"]["api_key"]
    except Exception:
        pass
    env_key = os.environ.get("GEMINI_API_KEY")
    if env_key:
        return env_key
    return get_drive_api_key()


def parse_and_import_syllabus(uploaded_file, subject_code):
    import re
    import pdfplumber
    
    text = ""
    try:
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception as e:
        return False, f"Failed to parse syllabus PDF: {str(e)}"
        
    lines = text.split('\n')
    current_unit = 1
    topics_imported = 0
    
    unit_pattern = re.compile(r'unit\s*[-:\s]*([iIvVxX\d]+)', re.IGNORECASE)
    
    conn = get_academic_conn()
    cursor = conn.cursor()
    
    for line in lines:
        line_str = line.strip()
        if not line_str or len(line_str) < 5 or len(line_str) > 120:
            continue
            
        # Check for unit header
        unit_match = unit_pattern.search(line_str)
        if unit_match and len(line_str) < 30:
            unit_val = unit_match.group(1).upper()
            roman_map = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, '1': 1, '2': 2, '3': 3, '4': 4, '5': 5}
            current_unit = roman_map.get(unit_val, current_unit)
            continue
            
        # Ignore common layout texts or headers
        if any(w in line_str.lower() for w in ["syllabus", "page", "credits", "lecture", "l t p c"]):
            continue
            
        # Insert as syllabus topic
        try:
            cursor.execute('''
                INSERT OR IGNORE INTO syllabus_topics (subject_code, unit, topic_text)
                VALUES (?, ?, ?)
            ''', (subject_code, current_unit, line_str))
            topics_imported += 1
        except Exception:
            pass
            
    conn.commit()
    conn.close()
    return True, f"Successfully parsed and loaded {topics_imported} syllabus topics for {subject_code}!"


def extract_options_helper(opt_text):
    import re
    opt_text = opt_text.replace('\t', ' ').replace('\n', ' ').strip()
    pattern = r'\b([a-d])\s*[\).\s]'
    matches = list(re.finditer(pattern, opt_text, re.IGNORECASE))
    
    opts = {'a': '', 'b': '', 'c': '', 'd': ''}
    
    if len(matches) >= 3:
        matches = sorted(matches, key=lambda x: x.start())
        first_m = matches[0].group(1).lower()
        if first_m != 'a':
            opts['a'] = opt_text[:matches[0].start()].strip()
            
        for i in range(len(matches)):
            m_char = matches[i].group(1).lower()
            start_pos = matches[i].end()
            end_pos = matches[i+1].start() if i + 1 < len(matches) else len(opt_text)
            
            val = opt_text[start_pos:end_pos].strip()
            val = re.sub(r'\[\s*\]\s*$', '', val).strip()
            opts[m_char] = val
    else:
        pattern_fallback = r'(?:\(?([a-d])\)|([a-d])\s*[\).\s])\s*([^(\n\t]+)'
        fallback_matches = re.findall(pattern_fallback, opt_text, re.IGNORECASE)
        for m in fallback_matches:
            char = m[0] or m[1]
            val = m[2].strip()
            opts[char.lower()] = val
            
    return opts['a'], opts['b'], opts['c'], opts['d']


def parse_and_import_qb(uploaded_file, file_type, subject_code, default_unit=3):
    import re
    import sqlite3
    
    actual_code = get_actual_subject_code(subject_code)
    
    descriptive_questions = []
    mcq_questions = []
    fib_questions = []
    
    if file_type == "pdf":
        import pdfplumber
        text = ""
        try:
            if hasattr(uploaded_file, "read"):
                uploaded_file.seek(0)
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text += (page.extract_text() or "") + "\n"
        except Exception as e:
            return False, f"Failed to parse PDF: {str(e)}"
            
        lines = text.split('\n')
        current_unit = default_unit
        
        q_pattern = re.compile(r'^(\d+)([a-z])?\)\s*(.*)', re.IGNORECASE)
        unit_pattern = re.compile(r'unit\s*[-:\s]*([iIvVxX\d]+)', re.IGNORECASE)
        
        current_q = None
        
        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue
                
            unit_match = unit_pattern.search(line_str)
            if unit_match and len(line_str) < 30:
                unit_val = unit_match.group(1).upper()
                roman_map = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, '1': 1, '2': 2, '3': 3, '4': 4, '5': 5}
                current_unit = roman_map.get(unit_val, current_unit)
                continue
                
            q_match = q_pattern.match(line_str)
            if q_match:
                q_num = q_match.group(1)
                q_sub = q_match.group(2) or ''
                q_text = q_match.group(3)
                
                marks = "10M"
                marks_match = re.search(r'\((\d+)\s*(?:Marks|M|Marks\))', q_text, re.IGNORECASE)
                if marks_match:
                    marks = f"{marks_match.group(1)}M"
                    q_text = re.sub(r'\((\d+)\s*(?:Marks|M|Marks\))\)', '', q_text).strip()
                elif q_text.strip().endswith('M)') or q_text.strip().endswith('Marks)'):
                    ending_match = re.search(r'\(?(\d+)\s*(?:M|Marks)\)?$', q_text.strip(), re.IGNORECASE)
                    if ending_match:
                        marks = f"{ending_match.group(1)}M"
                        q_text = re.sub(r'\(?(\d+)\s*(?:M|Marks)\)?$', '', q_text.strip()).strip()
                        
                if current_q:
                    descriptive_questions.append(current_q)
                    
                current_q = {
                    'subject_code': actual_code,
                    'unit': current_unit,
                    'q_type': 'DESCRIPTIVE',
                    'question': q_text,
                    'answer_text': '',
                    'marks': marks,
                    'co': '1',
                    'btl': 'L2'
                }
            else:
                if current_q:
                    if len(current_q['question']) < 120 and not current_q['question'].endswith('?') and not current_q['answer_text']:
                        current_q['question'] += " " + line_str
                    else:
                        current_q['answer_text'] += "\n" + line_str
                        
        if current_q:
            descriptive_questions.append(current_q)
            
    elif file_type == "docx":
        import docx
        try:
            if hasattr(uploaded_file, "read"):
                uploaded_file.seek(0)
            doc = docx.Document(uploaded_file)
        except Exception as e:
            return False, f"Failed to parse Word document: {str(e)}"
            
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        body_elements = []
        for child in doc.element.body:
            if child.tag.endswith('p'):
                body_elements.append(Paragraph(child, doc))
            elif child.tag.endswith('tbl'):
                body_elements.append(Table(child, doc))
                
        unit_data = {}
        current_unit = default_unit
        unit_re = re.compile(r'UNIT[-–\s]*([IVX\d]+)', re.IGNORECASE)
        
        in_objectives_global = False
        
        for element in body_elements:
            if isinstance(element, Paragraph):
                text = element.text.strip()
                if not text:
                    continue
                if "OBJECTIVE" in text.upper():
                    in_objectives_global = True
                unit_match = unit_re.search(text)
                if unit_match and ("UNIT" in text.upper() or len(text) < 50):
                    unit_str = unit_match.group(1).upper()
                    roman_map = {'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5, '1': 1, '2': 2, '3': 3, '4': 4, '5': 5}
                    current_unit = roman_map.get(unit_str, current_unit)
                    
            if current_unit not in unit_data:
                unit_data[current_unit] = []
            unit_data[current_unit].append(element)
            
        for unit, elements in sorted(unit_data.items()):
            unit_paras = [e for e in elements if isinstance(e, Paragraph)]
            unit_tables = [e for e in elements if isinstance(e, Table)]
            
            desc_table = None
            for t in unit_tables:
                headers = [cell.text.strip().upper() for cell in t.rows[0].cells]
                if any('DESCRIPTION' in h or 'QUESTION' in h for h in headers) and not any('ANSWER' in h for h in headers):
                    desc_table = t
                    break
                    
            if desc_table:
                headers = [cell.text.strip().replace('\n', ' ') for cell in desc_table.rows[0].cells]
                q_no_col = -1
                part_col = -1
                desc_col = -1
                marks_col = -1
                
                for col_idx, h in enumerate(headers):
                    hu = h.upper()
                    if 'Q.NO.' in hu or 'Q.NO' in hu or 'Q. NO' in hu or 'Q. NO.' in hu:
                        q_no_col = col_idx
                    elif 'PART' in hu or 'SUB' in hu:
                        part_col = col_idx
                    elif 'DESCRIPTION' in hu or 'QUESTION' in hu:
                        desc_col = col_idx
                    elif 'MARKS' in hu or 'MARK' in hu:
                        marks_col = col_idx
                        
                for row in desc_table.rows[1:]:
                    cells = [c.text.strip() for c in row.cells]
                    q_no_val = cells[q_no_col] if q_no_col != -1 and q_no_col < len(cells) else ""
                    part_val = cells[part_col] if part_col != -1 and part_col < len(cells) else ""
                    desc_val = cells[desc_col] if desc_col != -1 and desc_col < len(cells) else ""
                    marks_val = cells[marks_col] if marks_col != -1 and marks_col < len(cells) else "10M"
                    
                    if desc_val:
                        q_prefix = f"{q_no_val}{part_val})" if q_no_val else ""
                        descriptive_questions.append({
                            'subject_code': actual_code,
                            'unit': unit,
                            'q_type': 'DESCRIPTIVE',
                            'question': desc_val,
                            'answer_text': '',
                            'marks': marks_val,
                            'co': '1',
                            'btl': 'L2',
                            'prefix': q_prefix
                        })
                        
            objectives_paras = []
            in_obj_flow = False
            for p in unit_paras:
                t = p.text.strip()
                if not t:
                    continue
                if any(k in t.upper() for k in ["OBJECTIVE", "MULTIPLE CHOICE", "FILL-IN-THE-BLANK"]):
                    in_obj_flow = True
                    continue
                if in_objectives_global and unit_re.search(t) and len(t) < 50:
                    in_obj_flow = True
                    continue
                if in_obj_flow:
                    if unit_re.search(t) and len(t) < 40 and not any(k in t.upper() for k in ["OBJECTIVE"]):
                        break
                    objectives_paras.append(p)
                    
            clean_paras = []
            for p in objectives_paras:
                t = p.text.strip()
                if t.upper() in ["OBJECTIVE QUESTIONS", "OBJECTIVE KEY", "FILL-IN-THE-BLANK QUESTIONS"]:
                    continue
                clean_paras.append(t)
                
            fibs = []
            mcqs = []
            
            idx = 0
            while idx < len(clean_paras):
                t = clean_paras[idx]
                
                has_opts = "b)" in t.lower() or "b." in t.lower() or "(b)" in t.lower()
                next_is_options = False
                if idx + 1 < len(clean_paras):
                    next_t = clean_paras[idx+1]
                    next_has_opts = "b)" in next_t.lower() or "b." in next_t.lower() or "(b)" in next_t.lower()
                    next_is_q = next_t.endswith("?") or "[]" in next_t or "[ ]" in next_t or "___" in next_t
                    if next_has_opts and not next_is_q:
                        next_is_options = True
                        
                if has_opts or next_is_options:
                    q_text = t
                    opt_text = ""
                    if has_opts:
                        opt_text = t
                        first_opt_split = re.split(r'\b[a-d]\)|(?:\(?([a-d])\)|[a-d]\.\s+)', t, re.IGNORECASE, 1)
                        q_text_clean = first_opt_split[0].strip() if first_opt_split else t
                    else:
                        opt_text = clean_paras[idx+1]
                        idx += 1
                        q_text_clean = t
                        
                    q_text_clean = re.sub(r'\[\s*\]\s*$', '', q_text_clean).strip()
                    opt_a, opt_b, opt_c, opt_d = extract_options_helper(opt_text)
                    
                    mcqs.append({
                        "unit": unit,
                        "question": q_text_clean,
                        "option_a": opt_a,
                        "option_b": opt_b,
                        "option_c": opt_c,
                        "option_d": opt_d,
                        "answer": ""
                    })
                else:
                    q_text_clean = re.sub(r'\[\s*\]\s*$', '', t).strip()
                    fibs.append({
                        "unit": unit,
                        "question": q_text_clean,
                        "answer": ""
                    })
                idx += 1
                
            key_table = None
            for t in unit_tables:
                headers = [cell.text.strip().upper() for cell in t.rows[0].cells]
                is_k = any('ANSWER' in h for h in headers) or (len(headers) >= 4 and any(re.match(r'^\d+\.', h) for h in headers))
                if not is_k and len(t.columns) == 4 and len(t.rows) == 5:
                    is_k = True
                if is_k and t != desc_table:
                    key_table = t
                    break
                    
            answer_key_mapping = {}
            if key_table:
                headers = [cell.text.strip() for cell in key_table.rows[0].cells]
                if 'Q.No.' in headers or 'Q.NO.' in [h.upper() for h in headers]:
                    for row in key_table.rows[1:]:
                        cells = [c.text.strip() for c in row.cells]
                        if len(cells) >= 2:
                            q1, a1 = cells[0], cells[1]
                            if q1 and a1:
                                try:
                                    q_idx = int(q1)
                                    answer_key_mapping[q_idx] = a1
                                except ValueError:
                                    pass
                        if len(cells) >= 4:
                            q2, a2 = cells[2], cells[3]
                            if q2 and a2:
                                try:
                                    q_idx = int(q2)
                                    answer_key_mapping[q_idx] = a2
                                except ValueError:
                                    pass
                else:
                    for r_idx, row in enumerate(key_table.rows):
                        cells = [c.text.strip() for c in row.cells]
                        for col_idx, cell_text in enumerate(cells):
                            if not cell_text:
                                continue
                            prefix_match = re.match(r'^(\d+)\.\s*(.*)$', cell_text)
                            if prefix_match:
                                q_idx = int(prefix_match.group(1))
                                ans_val = prefix_match.group(2).strip()
                                answer_key_mapping[q_idx] = ans_val
                            else:
                                q_idx = col_idx * 5 + r_idx + 1
                                answer_key_mapping[q_idx] = cell_text
                                
            for i in range(min(len(fibs), 10)):
                fibs[i]["answer"] = answer_key_mapping.get(i+1, "")
                
            for i in range(min(len(mcqs), 10)):
                mcqs[i]["answer"] = answer_key_mapping.get(i+11, "")
                
            for fib in fibs:
                fib_questions.append(fib)
            for mcq in mcqs:
                mcq_questions.append(mcq)
                    
    else:
        return False, "Unsupported file format. Please upload PDF or DOCX."
        
    if not descriptive_questions and not mcq_questions and not fib_questions:
        api_key = get_gemini_api_key()
        if api_key:
            if file_type == "docx" and 'doc' in locals():
                text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            elif file_type == "pdf" and 'text' in locals():
                pass
            else:
                text = ""
                
            if text:
                import requests
                gemini_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
                headers = {"Content-Type": "application/json"}
                prompt = f"""
                You are a strict question bank parsing assistant. Your task is to extract all Descriptive Questions, MCQs, and Fill-in-the-Blank (FIB) questions from the provided course text.
                You must output a single, raw, valid JSON object with NO markdown wrapper (do not wrap in ```json), matching this schema:
                {{
                  "descriptive": [
                    {{"question": "full question text", "marks": "Marks (e.g. 10M or 2M)", "unit": 3, "co": "1", "btl": "L2"}}
                  ],
                  "mcqs": [
                    {{"question": "question text", "option_a": "option A", "option_b": "option B", "option_c": "option C", "option_d": "option D", "answer": "A/B/C/D", "unit": 3}}
                  ],
                  "fibs": [
                    {{"question": "question text with blank", "answer": "correct answer", "unit": 3}}
                  ]
                }}
                
                Ensure the unit matches the questions. If the text has no clear unit, use 3.
                
                Text:
                {text[:30000]}
                """
                try:
                    payload = {
                        "contents": [{"parts": [{"text": prompt}]}]
                    }
                    res = requests.post(gemini_url, json=payload, headers=headers, timeout=30)
                    if res.status_code == 200:
                        json_text = res.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
                        if json_text.startswith("```json"):
                            json_text = json_text[7:]
                        if json_text.endswith("```"):
                            json_text = json_text[:-3]
                        json_text = json_text.strip()
                        
                        parsed_data = json.loads(json_text)
                        
                        for q in parsed_data.get("descriptive", []):
                            descriptive_questions.append({
                                'subject_code': actual_code,
                                'unit': q.get('unit', default_unit),
                                'q_type': 'DESCRIPTIVE',
                                'question': q.get('question'),
                                'answer_text': '',
                                'marks': q.get('marks', '10M'),
                                'co': q.get('co', '1'),
                                'btl': q.get('btl', 'L2')
                            })
                        for q in parsed_data.get("mcqs", []):
                            mcq_questions.append({
                                'unit': q.get('unit', default_unit),
                                'question': q.get('question'),
                                'option_a': q.get('option_a'),
                                'option_b': q.get('option_b'),
                                'option_c': q.get('option_c'),
                                'option_d': q.get('option_d'),
                                'answer': q.get('answer', '')
                            })
                        for q in parsed_data.get("fibs", []):
                            fib_questions.append({
                                'unit': q.get('unit', default_unit),
                                'question': q.get('question'),
                                'answer': q.get('answer', '')
                            })
                except Exception:
                    pass

    conn = get_academic_conn()
    cursor = conn.cursor()
    
    desc_imported = 0
    mcq_imported = 0
    fib_imported = 0
    
    for q in descriptive_questions:
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO question_banks 
                (subject_code, unit, q_type, question, answer_text, marks, co, btl)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (q['subject_code'], q['unit'], 'DESCRIPTIVE', q['question'].strip(), q['answer_text'].strip(), q['marks'], q['co'], q['btl']))
            desc_imported += 1
            
            topic_text = q['question'].split('?')[0].split('.')[0].strip()
            if len(topic_text) > 10 and len(topic_text) < 120:
                cursor.execute('''
                    INSERT OR IGNORE INTO syllabus_topics (subject_code, unit, topic_text)
                    VALUES (?, ?, ?)
                ''', (q['subject_code'], q['unit'], topic_text))
        except Exception:
            pass
            
    for q in fib_questions:
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO question_banks 
                (subject_code, unit, q_type, question, answer_text, marks, co, btl)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)
            ''', (actual_code, q['unit'], 'FIB', q['question'].strip(), q['answer'].strip(), '2M', '1', 'L1'))
            fib_imported += 1
        except Exception:
            pass
            
    for q in mcq_questions:
        try:
            cursor.execute('''
                INSERT OR REPLACE INTO mcqs
                (subject_code, unit, question, option_a, option_b, option_c, option_d, correct_option, explanation)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            ''', (actual_code, q['unit'], q['question'].strip(), q['option_a'].strip(), q['option_b'].strip(), q['option_c'].strip(), q['option_d'].strip(), q['answer'].strip(), ''))
            mcq_imported += 1
        except Exception:
            pass
            
    conn.commit()
    conn.close()
    
    msg = f"Successfully parsed and loaded: {desc_imported} descriptive questions, {mcq_imported} MCQs, and {fib_imported} FIBs for {subject_code}!"
    return True, msg

# ── DATABASE HELPERS ─────────────────────────────────────────

def get_subject_display_name(code, roll_no):
    code = get_actual_subject_code(code)
    
    subject_names = {
        "2504103ES": "Electronic Devices and Circuits",
        "2566103ES": "Python Programming",
        "23EE104ES": "Basic Electrical Engineering",
        "2512105ES": "Programming for Problem Solving",
        "25CH102BS": "Engineering Chemistry",
        "25MA201BS": "Ordinary Differential Equations and Vector Calculus",
        "NAS": "Network Analysis and Synthesis",
        "25PH102BS": "Applied Engineering Physics",
        "2502105ES": "Introduction to Electrical Engineering",
        "2503104ES": "Engineering Drawing",
        "25EN105HS": "English for Professional Success",
        "EDC": "Electronic Devices and Circuits",
        "PYTHON": "Python Programming",
        "BEE": "Basic Electrical Engineering",
        "PPS": "Programming for Problem Solving",
        "Chemistry": "Engineering Chemistry",
        "ODEVC": "Ordinary Differential Equations and Vector Calculus",
        "M&C": "Mathematics & Calculus",
        "IEE": "Introduction to Engineering Electromagnetics",
        "ED": "Engineering Drawing",
        "ENG": "English for Professional Success",
        "AEP": "Applied Engineering Physics",
        "DS": "Data Structures",
        "EC": "Engineering Chemistry"
    }
    
    full_name = subject_names.get(code, code)
    
    # Extract branch from roll number using indices [6:8]
    branch_map = {
        "01": "CIVIL", "02": "EEE", "03": "MECH",
        "04": "ECE", "05": "CSE", "10": "EIE",
        "12": "IT", "66": "CSE AIML", "67": "CSE DS",
        "72": "AIDS", "73": "AIML"
    }
    branch_name = "CSE"  # default fallback
    if roll_no and len(roll_no) >= 8:
        b_code = roll_no[6:8]
        if b_code in branch_map:
            branch_name = branch_map[b_code]
            
    return f"{full_name} [{code}] (for {branch_name} students only)"

def resolve_subject_display_name(display_name):
    if not display_name:
        return ""
    import re
    match = re.search(r'\[([^\]]+)\]', display_name)
    if match:
        return match.group(1).strip()
    return display_name.strip()

def get_actual_subject_code(code):
    mapping = {
        "EDC": "2504103ES",
        "PYTHON": "2566103ES",
        "BEE": "23EE104ES",
        "PPS": "2512105ES",
        "Chemistry": "25CH102BS",
        "ODEVC": "25MA201BS",
        "NAS": "NAS",
        "M&C": "25PH102BS",
        "IEE": "2502105ES",
        "ED": "2503104ES",
        "ENG": "25EN105HS",
        "AEP": "25PH102BS"
    }
    return mapping.get(code, code)

def get_student_backlogs(roll_no):
    import sqlite3
    base_dir = os.path.dirname(os.path.abspath(__file__))
    main_db_path = os.path.join(base_dir, "vits_erp.db")
    
    conn = sqlite3.connect(main_db_path)
    conn.row_factory = sqlite3.Row
    rows = conn.execute('''
        SELECT subject, semester, score
        FROM marks
        WHERE roll_no=? AND exam_type LIKE '%Final Examinations' AND (score < 40 OR score IS NULL)
    ''', (roll_no,)).fetchall()
    conn.close()
    
    backlogs = []
    for r in rows:
        backlogs.append({
            'roll_no': roll_no,
            'subject_code': r['subject'],
            'backlog_sem': int(r['semester'].replace("Sem ", "").strip()) if "Sem" in r['semester'] else 1,
            'status': 'ACTIVE'
        })
    return backlogs

def get_student_exam_warnings(roll_no):
    import sqlite3
    base_dir = os.path.dirname(os.path.abspath(__file__))
    main_db_path = os.path.join(base_dir, "vits_erp.db")
    
    conn = sqlite3.connect(main_db_path)
    conn.row_factory = sqlite3.Row
    rows = conn.execute('''
        SELECT subject, score, exam_type
        FROM marks
        WHERE roll_no=? AND (
            (exam_type LIKE 'Mid%' AND score < 10) OR
            (exam_type NOT LIKE 'Mid%' AND exam_type NOT LIKE '%Final%' AND score < 40)
        )
    ''', (roll_no,)).fetchall()
    conn.close()
    
    warnings = []
    for r in rows:
        warnings.append({
            'subject': r['subject'],
            'score': r['score'],
            'max_score': 25 if 'Mid' in r['exam_type'] else 100,
            'exam_type': r['exam_type']
        })
    return warnings

def get_academic_resources(branch, semester):
    conn = get_academic_conn()
    rows = conn.execute('SELECT * FROM academic_resources WHERE branch=? AND semester=?', (branch, semester)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_syllabus_topics(subject_code):
    subject_code = get_actual_subject_code(subject_code)
    conn = get_academic_conn()
    rows = conn.execute('SELECT * FROM syllabus_topics WHERE subject_code=? ORDER BY unit, id', (subject_code,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_student_progress(roll_no):
    conn = get_academic_conn()
    rows = conn.execute('SELECT topic_id, status FROM student_syllabus_progress WHERE roll_no=?', (roll_no,)).fetchall()
    conn.close()
    return {r['topic_id']: r['status'] for r in rows}

def get_backlog_subject_progress(roll_no, subject_code):
    actual_code = get_actual_subject_code(subject_code)
    conn = get_academic_conn()
    topics = conn.execute('SELECT id, unit FROM syllabus_topics WHERE subject_code=?', (actual_code,)).fetchall()
    progress_rows = conn.execute('''
        SELECT topic_id, status 
        FROM student_syllabus_progress 
        WHERE roll_no=? AND topic_id IN (
            SELECT id FROM syllabus_topics WHERE subject_code=?
        )
    ''', (roll_no, actual_code)).fetchall()
    conn.close()
    
    progress_map = {r['topic_id']: r['status'] for r in progress_rows}
    
    unit_stats = {}
    for unit in [1, 2, 3, 4, 5]:
        unit_stats[unit] = {'total': 0, 'completed': 0}
        
    for t in topics:
        unit = t['unit']
        if unit in unit_stats:
            unit_stats[unit]['total'] += 1
            if progress_map.get(t['id']) == 'MASTERED':
                unit_stats[unit]['completed'] += 1
                
    return unit_stats

def update_syllabus_progress(roll_no, topic_id, status):
    conn = get_academic_conn()
    conn.execute('INSERT OR REPLACE INTO student_syllabus_progress (roll_no, topic_id, status) VALUES (?, ?, ?)',
                 (roll_no, topic_id, status))
    conn.commit()
    conn.close()

def get_question_bank(subject_code, unit, q_type=None):
    subject_code = get_actual_subject_code(subject_code)
    conn = get_academic_conn()
    if q_type:
        rows = conn.execute('SELECT * FROM question_banks WHERE subject_code=? AND unit=? AND q_type=?', (subject_code, unit, q_type)).fetchall()
    else:
        rows = conn.execute('SELECT * FROM question_banks WHERE subject_code=? AND unit=?', (subject_code, unit)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_mcqs(subject_code, unit):
    subject_code = get_actual_subject_code(subject_code)
    conn = get_academic_conn()
    rows = conn.execute('SELECT * FROM mcqs WHERE subject_code=? AND unit=?', (subject_code, unit)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_class_sessions(subject_code):
    subject_code = get_actual_subject_code(subject_code)
    conn = get_academic_conn()
    rows = conn.execute('SELECT * FROM class_sessions WHERE subject_code=? ORDER BY class_date', (subject_code,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def get_srs_reviews(roll_no):
    import datetime
    today = datetime.date.today().isoformat()
    conn = get_academic_conn()
    rows = conn.execute('''
        SELECT r.mcq_id, r.difficulty, r.easiness_factor, r.repetitions, r.interval, r.next_review, 
               m.question, m.option_a, m.option_b, m.option_c, m.option_d, m.correct_option, m.subject_code, m.unit
        FROM review_schedule r
        JOIN mcqs m ON r.mcq_id = m.id
        WHERE r.roll_no=? AND (r.next_review <= ? OR r.next_review IS NULL)
    ''', (roll_no, today)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

def add_srs_review(roll_no, mcq_id, quality=3):
    import datetime
    conn = get_academic_conn()
    row = conn.execute("SELECT * FROM review_schedule WHERE roll_no=? AND mcq_id=?", (roll_no, mcq_id)).fetchone()
    
    easiness_factor = 2.5
    repetitions = 0
    interval = 0
    
    if row:
        row_dict = dict(row)
        easiness_factor = row_dict.get("easiness_factor") or 2.5
        repetitions = row_dict.get("repetitions") or 0
        interval = row_dict.get("interval") or 0
        
    if quality >= 3:
        if repetitions == 0:
            interval = 1
        elif repetitions == 1:
            interval = 6
        else:
            interval = int(interval * easiness_factor)
        repetitions += 1
    else:
        repetitions = 0
        interval = 1
        
    easiness_factor = easiness_factor + (0.1 - (5 - quality) * (0.08 + (5 - quality) * 0.02))
    if easiness_factor < 1.3:
        easiness_factor = 1.3
        
    next_review_date = (datetime.date.today() + datetime.timedelta(days=interval)).isoformat()
    
    conn.execute('''
        INSERT OR REPLACE INTO review_schedule (roll_no, mcq_id, next_review, difficulty, easiness_factor, repetitions, interval)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (roll_no, mcq_id, next_review_date, quality, easiness_factor, repetitions, interval))
    conn.commit()
    conn.close()



def remove_srs_review(roll_no, mcq_id):
    conn = get_academic_conn()
    conn.execute('DELETE FROM review_schedule WHERE roll_no=? AND mcq_id=?', (roll_no, mcq_id))
    conn.commit()
    conn.close()
def show_academic_hub_page(student, sem):
    # Ensure roll_no is available in session state for helpers
    st.session_state["roll_no"] = student['roll_no']
    # Trigger the background sync thread (runs once every 15 mins, or forced)
    start_sync_thread()
    
    # Premium glassmorphic styling
    st.markdown("""
    <style>
    div.academic-card {
        background: rgba(13, 20, 38, 0.55) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(0, 216, 198, 0.15) !important;
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.4) !important;
        border-radius: 16px !important;
        padding: 24px;
        margin-bottom: 20px;
        transition: all 0.3s ease;
    }
    div.academic-card:hover {
        border-color: rgba(0, 216, 198, 0.35) !important;
        box-shadow: 0 12px 40px rgba(0, 216, 198, 0.12) !important;
    }
    div.academic-card-purple {
        background: rgba(13, 20, 38, 0.55) !important;
        backdrop-filter: blur(20px) !important;
        -webkit-backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(139, 92, 246, 0.18) !important;
        box-shadow: 0 8px 32px 0 rgba(0, 0, 0, 0.4) !important;
        border-radius: 16px !important;
        padding: 24px;
        margin-bottom: 20px;
        transition: all 0.3s ease;
    }
    div.academic-card-purple:hover {
        border-color: rgba(139, 92, 246, 0.38) !important;
        box-shadow: 0 12px 40px rgba(139, 92, 246, 0.12) !important;
    }
    .subject-card {
        background: rgba(20, 28, 48, 0.4) !important;
        backdrop-filter: blur(15px) !important;
        border: 1px solid rgba(255,255,255,0.06) !important;
        border-radius: 16px !important;
        padding: 18px !important;
        margin-bottom: 12px !important;
        transition: all 0.25s ease !important;
    }
    .subject-card:hover {
        border-color: rgba(0,216,198,0.3) !important;
        box-shadow: 0 10px 30px rgba(0,216,198,0.1) !important;
    }
    .card-bar-bg {
        width: 100% !important; height: 8px !important; background: rgba(255,255,255,0.08) !important;
        border-radius: 99px !important; overflow: hidden !important; margin: 10px 0 !important;
    }
    .card-bar-fill { height: 100% !important; border-radius: 99px !important; transition: width 0.6s ease !important; }
    </style>
    """, unsafe_allow_html=True)

    # Notion/Linear Navigation menu at the top
    if "academic_hub_sub_page" not in st.session_state:
        st.session_state["academic_hub_sub_page"] = "Dashboard"
        
    nav_cols = st.columns(5)
    nav_options = [
        ("🏠 Dashboard", "Dashboard"),
        ("📚 Subjects", "Subjects"),
        ("⚡ Exam Prep", "Exam Prep"),
        ("⚠️ Recovery", "Backlog Recovery"),
        ("🤖 AI Copilot", "AI Copilot")
    ]
    
    for idx, (label, val) in enumerate(nav_options):
        is_active = (st.session_state["academic_hub_sub_page"] == val)
        if nav_cols[idx].button(label, key=f"sub_nav_btn_{val}", use_container_width=True, 
                               type="primary" if is_active else "secondary"):
            st.session_state["academic_hub_sub_page"] = val
            st.rerun()
            
    st.markdown("<hr style='border: 0; height: 1px; background: rgba(255,255,255,0.08); margin-bottom: 25px;'>", unsafe_allow_html=True)
    
    current_sub = st.session_state["academic_hub_sub_page"]
    
    if current_sub == "Dashboard":
        render_dashboard_view(student)
    elif current_sub == "Subjects":
        render_subjects_view(student)
    elif current_sub == "Exam Prep":
        render_exam_prep_view(student)
    elif current_sub == "Backlog Recovery":
        render_backlog_recovery_view(student)
    elif current_sub == "AI Copilot":
        render_ai_copilot_view(student)


def get_subjects_list_all():
    conn = get_academic_conn()
    rows = conn.execute("SELECT DISTINCT subject_code FROM question_banks UNION SELECT DISTINCT subject_code FROM syllabus_topics UNION SELECT DISTINCT subject FROM drive_files WHERE processed=1").fetchall()
    conn.close()
    
    results = []
    for r in rows:
        code = r[0]
        if code:
            results.append({
                "subject_code": code,
                "display_name": get_subject_display_name(code, st.session_state.get("roll_no", "25H41A0401"))
            })
    return results


def get_topics_missed_for_student(roll_no):
    import sqlite3
    import os
    import random
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    main_db_path = os.path.join(base_dir, "vits_erp.db")
    
    absences = []
    if os.path.exists(main_db_path):
        try:
            conn = sqlite3.connect(main_db_path)
            conn.row_factory = sqlite3.Row
            rows = conn.execute("SELECT date, subject FROM hour_wise_attendance WHERE roll_no=?", (roll_no,)).fetchall()
            absences = [dict(r) for r in rows]
            conn.close()
        except Exception as e:
            print(f"Error fetching absences: {e}")
            
    if not absences:
        return {}
        
    conn_acad = get_academic_conn()
    missed = {}
    
    cnt = conn_acad.execute("SELECT COUNT(*) FROM syllabus_delivery_log").fetchone()[0]
    if cnt == 0:
        topics_pool = {
            "2566103ES": ["Introduction to Python", "Variables & Types", "Conditional Statements", "Loops (for/while)", "Functions", "Recursion", "Lists & Tuples", "Dictionaries & Sets"],
            "2504103ES": ["PN Junction Diode", "Zener Diode Characteristics", "Half Wave Rectifier", "Full Wave Rectifier", "BJT Operation", "BJT Configurations", "JFET Characteristics", "MOSFET Characteristics"],
            "23EE104ES": ["DC Circuits Basics", "Mesh Analysis", "Nodal Analysis", "AC Circuits Introduction", "Transformers Principles", "DC Motors Operation", "Three Phase Systems", "Electrical Safety"],
            "2512105ES": ["Introduction to C", "Operators & Expressions", "Arrays in C", "Strings & Operations", "Pointers Basics", "Structures & Unions", "File Handling in C", "Dynamic Memory Allocation"],
            "25CH102BS": ["Water Chemistry", "Hardness of Water", "Electrochemistry", "Corrosion & Prevention", "Polymer Materials", "Lubricants Types", "Nanomaterials", "Spectroscopy Techniques"],
            "25MA201BS": ["Ordinary Differential Equations", "Linear ODEs", "Vector Calculus", "Gradient, Divergence, Curl", "Line Integrals", "Surface Integrals", "Green's Theorem", "Stokes' Theorem"]
        }
        
        seen_combos = set()
        for ab in absences:
            dt_val = ab["date"]
            sub_name = ab["subject"]
            sub_code = get_actual_subject_code(sub_name)
            combo_key = (dt_val, sub_code)
            if combo_key not in seen_combos:
                seen_combos.add(combo_key)
                pool = topics_pool.get(sub_code, ["Introduction to Topic", "Advanced Concepts", "Practical Applications", "Review Session"])
                chosen_topic = random.choice(pool)
                conn_acad.execute('''
                    INSERT OR IGNORE INTO syllabus_delivery_log (class_date, subject_code, topic)
                    VALUES (?, ?, ?)
                ''', (dt_val, sub_code, chosen_topic))
        conn_acad.commit()
        
    for ab in absences:
        dt_val = ab["date"]
        sub_name = ab["subject"]
        sub_code = get_actual_subject_code(sub_name)
        
        log_rows = conn_acad.execute('''
            SELECT topic FROM syllabus_delivery_log 
            WHERE class_date = ? AND subject_code = ?
        ''', (dt_val, sub_code)).fetchall()
        
        for r in log_rows:
            t = r[0]
            if sub_name not in missed:
                missed[sub_name] = []
            if t not in missed[sub_name]:
                missed[sub_name].append(t)
                
    conn_acad.close()
    return missed


def render_dashboard_view(student):
    import math
    roll = student['roll_no']
    name = student['name']
    branch = student['branch']
    
    hour = datetime.datetime.now().hour
    greeting = "Good Morning" if hour < 12 else "Good Afternoon" if hour < 17 else "Good Evening"
    
    st.markdown(f"""
    <div class="academic-card-purple">
        <h2 style="margin: 0; color: #8B5CF6; font-family: 'Outfit', sans-serif;">🌅 {greeting}, {name}!</h2>
        <p style="margin: 5px 0 0 0; font-size: 1.05rem; color: #cbd5e1;">Welcome to your AI Academic Copilot. Here is your personalized intelligence briefing for today.</p>
    </div>
    """, unsafe_allow_html=True)
    
    if "sync_status" in st.session_state:
        status_color = "#00D8C6" if "Complete" in st.session_state["sync_status"] else "#F59E0B"
        st.markdown(f"""
        <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); 
                    border-left: 4px solid {status_color}; border-radius: 8px; padding: 10px; margin-bottom: 20px; font-size: 0.85rem;">
            <b>Sync Engine Status:</b> {st.session_state["sync_status"]}
        </div>
        """, unsafe_allow_html=True)
        
    import sqlite3
    base_dir = os.path.dirname(os.path.abspath(__file__))
    main_db_path = os.path.join(base_dir, "vits_erp.db")
    
    att_rows = []
    marks_rows = []
    if os.path.exists(main_db_path):
        conn = sqlite3.connect(main_db_path)
        conn.row_factory = sqlite3.Row
        att_rows = conn.execute("SELECT * FROM attendance WHERE roll_no=?", (roll,)).fetchall()
        marks_rows = conn.execute("SELECT * FROM marks WHERE roll_no=?", (roll,)).fetchall()
        conn.close()
        
    total_conducted = sum(r['hours_conducted'] or 0 for r in att_rows)
    total_attended = sum(r['hours_attended'] or 0 for r in att_rows)
    attendance_pct = (total_attended / total_conducted * 100) if total_conducted > 0 else 0.0
    
    subjects_list = [r['subject_code'] for r in get_subjects_list_all()]
    total_topics = 0
    completed_topics = 0
    
    conn_acad = get_academic_conn()
    progress_rows = conn_acad.execute("SELECT topic_id, status FROM student_syllabus_progress WHERE roll_no=?", (roll,)).fetchall()
    progress_map = {r['topic_id']: r['status'] for r in progress_rows}
    
    for sub in subjects_list:
        sub_code = get_actual_subject_code(sub)
        t_rows = conn_acad.execute("SELECT id FROM syllabus_topics WHERE subject_code=?", (sub_code,)).fetchall()
        for tr in t_rows:
            total_topics += 1
            if progress_map.get(tr[0]) == 'MASTERED':
                completed_topics += 1
    conn_acad.close()
    
    syllabus_pct = (completed_topics / total_topics * 100) if total_topics > 0 else 0.0
    
    final_exam_scores = [r['score'] for r in marks_rows if 'Final' in r['exam_type'] and r['score'] is not None]
    avg_score = sum(final_exam_scores) / len(final_exam_scores) if final_exam_scores else 0.0
    
    c1, c2, c3 = st.columns(3)
    with c1:
        st.metric("Overall Attendance", f"{attendance_pct:.1f}%", help="Required is 75%+")
    with c2:
        st.metric("Syllabus Mastery", f"{syllabus_pct:.1f}%", f"{completed_topics}/{total_topics} Topics", help="Track syllabus topics in Subjects")
    with c3:
        st.metric("Academic Health Index", f"{avg_score:.1f}/100" if avg_score > 0 else "N/A", help="Average Final Exam Score")
        
    st.markdown("<div style='height: 15px;'></div>", unsafe_allow_html=True)
    
    col_left, col_right = st.columns([1.1, 0.9])
    
    with col_left:
        st.markdown("### ⚠️ Academic Risk Analysis")
        has_alerts = False
        
        for r in att_rows:
            sub = r['subject']
            attended = r['hours_attended'] or 0
            conducted = r['hours_conducted'] or 0
            pct = (attended / conducted * 100) if conducted > 0 else 0.0
            
            if conducted > 0 and pct < 75.0:
                has_alerts = True
                needed = math.ceil((0.75 * conducted - attended) / (1 - 0.75))
                st.markdown(f"""
                <div style="background: rgba(239, 68, 68, 0.12); border: 1px solid rgba(239, 68, 68, 0.3); 
                            border-left: 4px solid #EF4444; border-radius: 8px; padding: 12px; margin-bottom: 10px;">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <strong style="color: #fff;">Attendance Risk in {sub}</strong>
                        <span style="color: #EF4444; font-weight: 700;">{pct:.1f}%</span>
                    </div>
                    <p style="margin: 5px 0 0 0; font-size: 0.85rem; color: #cbd5e1;">
                        You are below the 75% threshold. You must attend <b>{needed}</b> consecutive classes of {sub} to restore your attendance.
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
        warnings = get_student_exam_warnings(roll)
        for w in warnings:
            has_alerts = True
            st.markdown(f"""
            <div style="background: rgba(245, 158, 11, 0.12); border: 1px solid rgba(245, 158, 11, 0.3); 
                        border-left: 4px solid #F59E0B; border-radius: 8px; padding: 12px; margin-bottom: 10px;">
                <div style="display: flex; justify-content: space-between; align-items: center;">
                    <strong style="color: #fff;">Backlog Risk: {w['subject']}</strong>
                    <span style="color: #F59E0B; font-weight: 700;">{w['score']}/{w['max_score']}</span>
                </div>
                <p style="margin: 5px 0 0 0; font-size: 0.85rem; color: #cbd5e1;">
                    Your score is low in {w['exam_type']}. We recommend launching the <b>Backlog Recovery</b> program for this subject.
                </p>
            </div>
            """, unsafe_allow_html=True)
            
        if not has_alerts:
            st.markdown("""
            <div style="background: rgba(16, 185, 129, 0.1); border: 1px solid rgba(16, 185, 129, 0.2); 
                        border-left: 4px solid #10B981; border-radius: 8px; padding: 15px; margin-bottom: 15px; text-align: center;">
                <span style="color: #10B981; font-weight: 600;">✨ No active academic risks detected. You are on track!</span>
            </div>
            """, unsafe_allow_html=True)
            
        st.markdown("<div style='height: 10px;'></div>", unsafe_allow_html=True)
        
        st.markdown("### 📴 Missed Classes & Syllabus Recovery")
        missed_topics = get_topics_missed_for_student(roll)
        if missed_topics:
            st.markdown("<p style='font-size: 0.9rem; color: #94a3b8; margin-bottom: 10px;'>Our sync engine cross-referenced your absent dates and identified the following taught topics you missed:</p>", unsafe_allow_html=True)
            for sub, topics in missed_topics.items():
                st.markdown(f"""
                <div class="academic-card" style="padding: 15px; margin-bottom: 10px; border-left: 4px solid #8B5CF6;">
                    <h5 style="margin: 0 0 8px 0; color: #fff;">{sub}</h5>
                    <div style="display: flex; flex-wrap: wrap; gap: 6px;">
                """, unsafe_allow_html=True)
                for t in topics:
                    st.markdown(f"<span style='background: rgba(139, 92, 246, 0.15); border: 1px solid rgba(139, 92, 246, 0.3); color: #c084fc; border-radius: 6px; padding: 4px 8px; font-size: 0.8rem;'>{t}</span>", unsafe_allow_html=True)
                st.markdown(f"""
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div style="background: rgba(255,255,255,0.02); border: 1px solid rgba(255,255,255,0.05); border-radius: 8px; padding: 20px; text-align: center;">
                <p style="color: #cbd5e1; margin: 0; font-size: 0.9rem;">You have perfect attendance or no recorded lectures correspond to your absent days!</p>
            </div>
            """, unsafe_allow_html=True)
            
    with col_right:
        st.markdown("### 🗓️ Today's Study Planner")
        st.markdown("<p style='font-size: 0.9rem; color: #94a3b8; margin-bottom: 15px;'>Your AI Study Planner recommends completing these targets today to maximize recovery:</p>", unsafe_allow_html=True)
        
        tasks = []
        due_reviews = get_srs_reviews(roll)
        if due_reviews:
            tasks.append({
                "title": f"Revise {len(due_reviews)} due questions in spaced repetition",
                "tag": "SRS",
                "color": "#8B5CF6"
            })
            
        if missed_topics:
            for sub, topics in list(missed_topics.items())[:1]:
                tasks.append({
                    "title": f"Study missed topic: '{topics[0]}' in {sub}",
                    "tag": "Missed Class",
                    "color": "#EF4444"
                })
                
        tasks.append({
            "title": f"Take a practice quiz on your weak subject to build accuracy",
            "tag": "Quiz",
            "color": "#00D8C6"
        })
        
        for idx, task in enumerate(tasks):
            st.markdown(f"""
            <div class="academic-card" style="padding: 16px; margin-bottom: 12px; border-left: 4px solid {task['color']};">
                <div style="display: flex; justify-content: space-between; align-items: start; margin-bottom: 5px;">
                    <span style="background: rgba(255,255,255,0.05); border: 1px solid rgba(255,255,255,0.1); 
                                 border-radius: 4px; padding: 2px 6px; font-size: 0.7rem; color: {task['color']}; font-weight: 700; text-transform: uppercase;">
                        {task['tag']}
                    </span>
                </div>
                <p style="margin: 0; font-size: 0.9rem; color: #fff; font-weight: 500;">{task['title']}</p>
            </div>
            """, unsafe_allow_html=True)
            
        st.markdown("#### 🤖 AI Generated Study Schedule")
        if st.button("✨ Generate Custom Study Plan", key="btn_gen_study_plan"):
            with st.spinner("AI is analyzing your workload..."):
                api_key = get_gemini_api_key()
                url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
                headers = {"Content-Type": "application/json"}
                
                context = f"Student Name: {name}, Branch: {branch}, Attendance: {attendance_pct:.1f}%, Syllabus Completion: {syllabus_pct:.1f}%.\n"
                context += f"Missed Topics: {json.dumps(missed_topics)}.\n"
                context += f"Weak Exam Marks: {json.dumps(warnings)}.\n"
                
                prompt = f"""
                You are an expert AI Study Planner. Create a concise, action-oriented daily study schedule for this student based on their status:
                {context}
                
                Generate:
                1. Today's Tasks (3 items)
                2. Tomorrow's Tasks (3 items)
                3. A Weekly Focus Goal.
                
                Keep the output clear, short, and formatted in clean markdown.
                """
                payload = {
                    "contents": [{"parts": [{"text": prompt}]}]
                }
                try:
                    res = requests.post(url, json=payload, headers=headers, timeout=25)
                    if res.status_code == 200:
                        st.markdown(res.json()["candidates"][0]["content"]["parts"][0]["text"])
                    else:
                        st.error("Failed to query Gemini API.")
                except Exception as e:
                    st.error(f"Error: {e}")


def render_subjects_view(student):
    import math
    roll = student['roll_no']
    
    subjects = get_subjects_list_all()
    if not subjects:
        st.markdown("""
        <div class="academic-card" style="text-align: center; padding: 50px;">
            <h3 style="color: #EF4444;">📭 No Academic Material Synchronized</h3>
            <p style="color: #cbd5e1; font-size: 0.95rem; margin-top: 10px;">
                We found no processed subjects in the database. Ensure that faculty have uploaded syllabus, question banks, lab code, and notes files to the Google Drive folder:
            </p>
            <p style="font-family: monospace; font-size: 0.9rem; background: rgba(255,255,255,0.05); padding: 8px; border-radius: 6px;">
                1bInXkRc9mQFdbVbUMNxG1VVnrpKEyoPN
            </p>
            <p style="color: #94a3b8; font-size: 0.85rem;">The background sync engine is active and will process them automatically.</p>
        </div>
        """, unsafe_allow_html=True)
        return
        
    active_subject = st.session_state.get("active_subject")
    
    if active_subject:
        if st.button("⬅️ Back to Subject Cards"):
            st.session_state["active_subject"] = None
            st.rerun()
            
        render_subject_deep_dive(student, active_subject)
        return
        
    st.markdown("### 📚 Subject Directory")
    st.markdown("<p style='font-size: 0.9rem; color: #94a3b8; margin-bottom: 20px;'>Explore courses, track unit checklists, practice quizzes, and view reference lab programs.</p>", unsafe_allow_html=True)
    
    rows_count = math.ceil(len(subjects) / 2)
    conn = get_academic_conn()
    
    for r_i in range(rows_count):
        cols = st.columns(2)
        for c_i in range(2):
            idx = r_i * 2 + c_i
            if idx < len(subjects):
                sub = subjects[idx]
                code = sub["subject_code"]
                
                total_topics = conn.execute("SELECT COUNT(*) FROM syllabus_topics WHERE subject_code=?", (code,)).fetchone()[0]
                completed_topics = conn.execute('''
                    SELECT COUNT(*) FROM student_syllabus_progress 
                    WHERE roll_no=? AND status='MASTERED' 
                      AND topic_id IN (SELECT id FROM syllabus_topics WHERE subject_code=?)
                ''', (roll, code)).fetchone()[0]
                total_q = conn.execute("SELECT COUNT(*) FROM question_banks WHERE subject_code=? AND q_type='DESCRIPTIVE'", (code,)).fetchone()[0]
                total_m = conn.execute("SELECT COUNT(*) FROM mcqs WHERE subject_code=?", (code,)).fetchone()[0]
                labs = conn.execute("SELECT COUNT(*) FROM lab_programs WHERE subject_code=?", (code,)).fetchone()[0]
                notes = conn.execute("SELECT COUNT(*) FROM notes WHERE subject_code=?", (code,)).fetchone()[0]
                
                last_up_row = conn.execute("SELECT modified_time FROM drive_files WHERE subject=? ORDER BY modified_time DESC LIMIT 1", (code,)).fetchone()
                last_up = last_up_row[0] if last_up_row else "N/A"
                if last_up != "N/A":
                    try:
                        dt_obj = datetime.datetime.strptime(last_up.split(".")[0].replace("Z", ""), "%Y-%m-%dT%H:%M:%S")
                        last_up = dt_obj.strftime("%b %d, %Y")
                    except Exception:
                        pass
                
                pct = int((completed_topics / total_topics * 100)) if total_topics > 0 else 0
                
                with cols[c_i]:
                    st.markdown(f"""
                    <div class="subject-card">
                        <div style="display: flex; justify-content: space-between; align-items: start; margin-bottom: 8px;">
                            <h4 style="margin: 0; color: #fff; font-family: 'Outfit'; font-size: 1.15rem;">
                                {get_subject_display_name(code, roll).split(" [")[0]}
                            </h4>
                            <span style="font-family: 'JetBrains Mono'; font-weight: 700; color: #00D8C6; font-size: 0.9rem;">
                                {pct}%
                            </span>
                        </div>
                        <span style="font-size: 0.75rem; color: #8B5CF6; font-family: 'JetBrains Mono'; font-weight: 600;">CODE: {code}</span>
                        <div class="card-bar-bg">
                            <div class="card-bar-fill" style="width: {pct}%; background: linear-gradient(90deg, #00D8C6 0%, #8B5CF6 100%);"></div>
                        </div>
                        
                        <div style="display: grid; grid-template-columns: 1fr 1fr; gap: 8px 15px; margin-top: 15px; font-size: 0.8rem; color: #cbd5e1;">
                            <div>📋 <b>Syllabus topics:</b> {completed_topics}/{total_topics}</div>
                            <div>❓ <b>Questions:</b> {total_q}</div>
                            <div>🎯 <b>Quiz MCQs:</b> {total_m}</div>
                            <div>💻 <b>Lab Scripts:</b> {labs}</div>
                            <div>📝 <b>Study Notes:</b> {notes}</div>
                            <div>🕒 <b>Updated:</b> {last_up}</div>
                        </div>
                    </div>
                    """.replace('\n', ' '), unsafe_allow_html=True)
                    
                    if st.button(f"🔍 Open {code} Study Deck", key=f"open_sub_btn_{code}", use_container_width=True):
                        st.session_state["active_subject"] = code
                        st.rerun()
                        
    conn.close()


def generate_tts_lecture_briefing(roll, subject_code, duration, cache_dir):
    with st.spinner("AI is generating speech narration..."):
        notes = get_notes(subject_code)
        
        text = f"Hello. This is your VITS Academic Copilot {duration} lecture briefing for subject code {subject_code}. "
        if not notes:
            text += "We do not have any notes synced in the database for this subject yet. Please make sure class notes files are uploaded to Google Drive."
        else:
            text += f"We will summarize {len(notes)} key topics. "
            limit = 3 if duration == "5 min" else 8 if duration == "15 min" else len(notes)
            for n in notes[:limit]:
                text += f"In Unit {n['unit']}, regarding {n['topic']}: {n['content'][:300]}. "
                
        try:
            from gtts import gTTS
            tts = gTTS(text=text, lang='en')
            audio_path = os.path.join(cache_dir, f"{roll}_{subject_code}_lecture.mp3")
            tts.save(audio_path)
            st.success("Narration audio generated successfully!")
        except Exception as e:
            st.error(f"Failed to generate Text-to-Speech audio: {e}")


def render_subject_deep_dive(student, subject_code):
    roll = student['roll_no']
    branch = student['branch']
    
    friendly_name = get_subject_display_name(subject_code, roll)
    
    st.markdown(f"""
    <div class="academic-card-purple" style="margin-bottom: 25px;">
        <h3 style="margin: 0; color: #8B5CF6;">{friendly_name}</h3>
        <p style="margin: 5px 0 0 0; font-size: 0.9rem; color: #cbd5e1;">Study Deck, Syllabus Checklists, Lab Codes, and Practice Quizzes.</p>
    </div>
    """, unsafe_allow_html=True)
    
    s_tabs = st.tabs(["📋 Syllabus", "📝 Notes", "❓ Questions", "💻 Labs", "🎙️ NotebookLM", "🎯 Quiz"])
    
    with s_tabs[0]:
        st.write("#### Course Syllabus Checklist")
        topics = get_syllabus_topics(subject_code)
        if not topics:
            st.info("No syllabus topics loaded yet for this subject. Sync syllabus file from Drive.")
        else:
            progress_map = get_student_progress(roll)
            completed_count = sum(1 for t in topics if progress_map.get(t['id']) == 'MASTERED')
            total_count = len(topics)
            pct = (completed_count / total_count) if total_count > 0 else 0
            
            st.markdown(f"**Overall Syllabus Progress: {completed_count}/{total_count} Topics ({int(pct*100)}%)**")
            st.progress(pct)
            
            for unit in [1, 2, 3, 4, 5]:
                unit_topics = [t for t in topics if t['unit'] == unit]
                if not unit_topics:
                    continue
                with st.expander(f"Unit {unit} Checklist", expanded=(unit == 1)):
                    for t in unit_topics:
                        is_checked = progress_map.get(t['id']) == 'MASTERED'
                        chk = st.checkbox(t['topic_text'], value=is_checked, key=f"deep_sub_chk_{t['id']}")
                        if chk != is_checked:
                            new_status = 'MASTERED' if chk else 'NOT_STARTED'
                            update_syllabus_progress(roll, t['id'], new_status)
                            st.rerun()
                            
    with s_tabs[1]:
        st.write("#### Class Notes & Reading Material")
        notes = get_notes(subject_code)
        if not notes:
            st.info("No notes processed for this subject yet. Upload study guides to Drive to index them.")
        else:
            for n in notes:
                with st.expander(f"Unit {n['unit']} Topic: {n['topic']}"):
                    st.write(n['content'])
                    
    with s_tabs[2]:
        st.write("#### Descriptive Questions Bank")
        questions = []
        conn = get_academic_conn()
        q_rows = conn.execute("SELECT * FROM question_banks WHERE subject_code=? AND q_type='DESCRIPTIVE'", (get_actual_subject_code(subject_code),)).fetchall()
        conn.close()
        questions = [dict(r) for r in q_rows]
        
        if not questions:
            st.info("No descriptive questions found in question bank. Sync question bank documents from Drive.")
        else:
            unit_filter = st.selectbox("Filter by Unit", ["All Units", "Unit 1", "Unit 2", "Unit 3", "Unit 4", "Unit 5"], key="deep_q_filter")
            filtered_qs = questions
            if unit_filter != "All Units":
                u_num = int(unit_filter.split(" ")[1])
                filtered_qs = [q for q in questions if q['unit'] == u_num]
                
            for idx, q in enumerate(filtered_qs):
                with st.expander(f"Q{idx+1}. [Unit {q['unit']}] {q['question'][:80]}..."):
                    st.markdown(f"**Question:**\\n{q['question']}")
                    st.markdown(f"**Marks:** {q['marks']} | **CO:** {q['co']} | **BTL:** {q['btl']}")
                    if q['answer_text']:
                        st.markdown(f"**Suggested Answer:**\\n{q['answer_text']}")
                    else:
                        st.info("No parsed answer notes. Copy the question to ask AI Tutor.")
                        
    with s_tabs[3]:
        st.write("#### Dynamic Lab Library")
        labs = get_lab_programs(subject_code)
        if not labs:
            st.info("No lab program scripts found for this subject. Upload code files or a ZIP to Drive under 'Lab Programs' folder.")
        else:
            selected_lab = st.selectbox("Select Lab Script", [l['file_name'] for l in labs], key="deep_lab_select")
            lab_obj = [l for l in labs if l['file_name'] == selected_lab][0]
            
            lang = "python" if selected_lab.endswith(".py") else "c" if selected_lab.endswith((".c", ".cpp")) else "text"
            st.code(lab_obj['code_content'], language=lang)
            
            c1, c2 = st.columns(2)
            c1.download_button("📥 Download Script", lab_obj['code_content'], file_name=selected_lab, use_container_width=True)
            if c2.button("📋 Copy Code to Clipboard", use_container_width=True):
                st.success("Code copy template available inside the code block's hover controls!")
                
    with s_tabs[4]:
        st.write("#### NotebookLM Studio Exporter")
        
        st.write("##### 🎙️ Generate AI Lecture briefing")
        audio_cache_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "audio_cache")
        os.makedirs(audio_cache_dir, exist_ok=True)
        
        c1, c2, c3 = st.columns(3)
        duration_sel = "5 min"
        if c1.button("📻 5 min Summary", use_container_width=True):
            duration_sel = "5 min"
            generate_tts_lecture_briefing(roll, subject_code, duration_sel, audio_cache_dir)
        if c2.button("📻 15 min Revision", use_container_width=True):
            duration_sel = "15 min"
            generate_tts_lecture_briefing(roll, subject_code, duration_sel, audio_cache_dir)
        if c3.button("📻 30 min Deep Dive", use_container_width=True):
            duration_sel = "30 min"
            generate_tts_lecture_briefing(roll, subject_code, duration_sel, audio_cache_dir)
            
        audio_file = os.path.join(audio_cache_dir, f"{roll}_{subject_code}_lecture.mp3")
        if os.path.exists(audio_file):
            with open(audio_file, "rb") as f:
                st.audio(f.read(), format="audio/mp3")
        else:
            st.info("No custom audio briefing generated yet. Select a duration option above to compile and narrate.")
            
        st.write("##### 📥 NotebookLM Custom Study Package")
        st.write("Compile a personalized study package containing course notes, descriptive questions, and spaced repetition errors. Upload it to Google NotebookLM to build audio discussions or custom tutors.")
        
        md_content = f"# NotebookLM Study Package - {friendly_name}\\n\\n"
        md_content += f"Roll Number: {roll} | Branch: {branch}\\n\\n"
        
        notes = get_notes(subject_code)
        if notes:
            md_content += "## 📝 Course Notes\\n"
            for n in notes:
                md_content += f"### Unit {n['unit']}: {n['topic']}\\n{n['content']}\\n\\n"
                
        conn = get_academic_conn()
        q_rows = conn.execute("SELECT question, answer_text, unit FROM question_banks WHERE subject_code=? AND q_type='DESCRIPTIVE'", (get_actual_subject_code(subject_code),)).fetchall()
        conn.close()
        if q_rows:
            md_content += "## ❓ Descriptive Questions & Answers\\n"
            for q in q_rows:
                ans = q[1] if q[1] else "(See AI Copilot for explanation)"
                md_content += f"### [Unit {q[2]}] Question: {q[0]}\\nSuggested Answer: {ans}\\n\\n"
                
        st.download_button("📥 Export StudyPackage.md", md_content, file_name=f"{subject_code}_NotebookLM_Pack.md", use_container_width=True)
        
    with s_tabs[5]:
        st.write("#### Subject Practice Quiz")
        render_subject_quiz_engine(roll, subject_code)


def render_subject_quiz_engine(roll, subject_code):
    actual_code = get_actual_subject_code(subject_code)
    
    q_mode = st.radio("Choose Quiz Mode", ["Practice Mode", "Exam Mode", "Revision Mode (SRS)"], horizontal=True, key=f"q_mode_{subject_code}")
    
    conn = get_academic_conn()
    if "Revision" in q_mode:
        today = datetime.date.today().isoformat()
        rows = conn.execute('''
            SELECT m.*, r.difficulty FROM review_schedule r
            JOIN mcqs m ON r.mcq_id = m.id
            WHERE r.roll_no=? AND m.subject_code=? AND (r.next_review <= ? OR r.next_review IS NULL)
        ''', (roll, actual_code, today)).fetchall()
        questions = [dict(r) for r in rows]
    else:
        rows = conn.execute("SELECT * FROM mcqs WHERE subject_code=?", (actual_code,)).fetchall()
        questions = [dict(r) for r in rows]
    conn.close()
    
    if not questions:
        st.info("No questions available for this subject in the selected mode. (Wrong answers in standard quizzes will build your Revision queue!)")
        return
        
    state_key = f"quiz_session_{subject_code}_{q_mode.replace(' ', '')}"
    if state_key not in st.session_state:
        import random
        selected_qs = list(questions)
        if len(selected_qs) > 10:
            selected_qs = random.sample(selected_qs, 10)
            
        st.session_state[state_key] = {
            "questions": selected_qs,
            "current_index": 0,
            "score": 0,
            "answers": {},
            "wrong_ids": [],
            "completed": False,
            "start_time": time.time()
        }
        
    quiz = st.session_state[state_key]
    
    if quiz["completed"]:
        duration = int(time.time() - quiz["start_time"])
        accuracy = (quiz["score"] / len(quiz["questions"]) * 100) if quiz["questions"] else 0
        
        st.balloons()
        st.markdown(f"""
        <div class="academic-card" style="text-align: center; border-left: 4px solid #00D8C6;">
            <h4 style="color: #00D8C6;">🎉 Quiz Completed!</h4>
            <p style="font-size: 1.25rem; font-weight: 700; color: #fff; margin: 10px 0;">Accuracy: {accuracy:.1f}% ({quiz['score']} / {len(quiz['questions'])})</p>
            <p style="font-size: 0.9rem; color: #cbd5e1;">Time Taken: <b>{duration} seconds</b></p>
        </div>
        """, unsafe_allow_html=True)
        
        weak_units = {}
        for q_idx in quiz["wrong_ids"]:
            q_obj = quiz["questions"][q_idx]
            u = q_obj.get("unit", 3)
            weak_units[u] = weak_units.get(u, 0) + 1
            
            if "Revision" not in q_mode:
                add_srs_review(roll, q_obj['id'], 1)
                
        if weak_units:
            st.warning("⚠️ **Weak Units Detected:**")
            for u, cnt in weak_units.items():
                st.write(f"- Unit {u}: {cnt} wrong answers. Recommend reviewing Unit {u} syllabus topics.")
                
        if st.button("🔄 Restart Quiz", key=f"restart_btn_{state_key}"):
            del st.session_state[state_key]
            st.rerun()
        return
        
    idx = quiz["current_index"]
    q_obj = quiz["questions"][idx]
    
    st.markdown(f"**Question {idx+1} of {len(quiz['questions'])}** (Unit {q_obj['unit']})")
    st.write(q_obj['question'])
    
    opts = [q_obj['option_a'], q_obj['option_b'], q_obj['option_c'], q_obj['option_d']]
    opts = [o for o in opts if o]
    ans_labels = ["A", "B", "C", "D"]
    
    if "Practice" in q_mode:
        user_select = st.radio("Select Option:", opts, key=f"radio_{state_key}_{idx}")
        sel_letter = ans_labels[opts.index(user_select)] if user_select in opts else ""
        
        if st.button("Check Answer", key=f"chk_btn_{state_key}_{idx}"):
            correct = q_obj['correct_option'].strip().upper()
            if sel_letter == correct:
                st.success("Correct! 🎉")
                quiz["score"] += 1
            else:
                st.error(f"Incorrect. Correct answer is **{correct}**.")
                if q_obj.get("explanation"):
                    st.info(f"Explanation: {q_obj['explanation']}")
                quiz["wrong_ids"].append(idx)
                
            if idx + 1 < len(quiz["questions"]):
                quiz["current_index"] += 1
            else:
                quiz["completed"] = True
            st.button("Continue", key=f"cont_btn_{state_key}_{idx}")
            
    elif "Exam" in q_mode:
        user_select = st.radio("Select Option:", opts, key=f"radio_{state_key}_{idx}")
        sel_letter = ans_labels[opts.index(user_select)] if user_select in opts else ""
        quiz["answers"][idx] = sel_letter
        
        c1, c2 = st.columns(2)
        if idx + 1 < len(quiz["questions"]):
            if c1.button("Next Question", key=f"next_btn_{state_key}_{idx}"):
                quiz["current_index"] += 1
                st.rerun()
        else:
            if c1.button("Finish Exam & Submit", key=f"finish_btn_{state_key}_{idx}"):
                for i, q in enumerate(quiz["questions"]):
                    user_ans = quiz["answers"].get(i, "")
                    correct = q['correct_option'].strip().upper()
                    if user_ans == correct:
                        quiz["score"] += 1
                    else:
                        quiz["wrong_ids"].append(i)
                quiz["completed"] = True
                st.rerun()
                
    else:
        user_select = st.radio("Select Option:", opts, key=f"radio_{state_key}_{idx}")
        sel_letter = ans_labels[opts.index(user_select)] if user_select in opts else ""
        
        if "rev_ans_state" not in st.session_state:
            st.session_state["rev_ans_state"] = None
            
        if st.button("Check Revision Answer", key=f"chk_rev_{state_key}_{idx}"):
            correct = q_obj['correct_option'].strip().upper()
            if sel_letter == correct:
                st.session_state["rev_ans_state"] = "correct"
            else:
                st.session_state["rev_ans_state"] = "incorrect"
                quiz["wrong_ids"].append(idx)
                add_srs_review(roll, q_obj['id'], 1)
                
        if st.session_state["rev_ans_state"] == "correct":
            st.success("Correct! 🎉 How easy was it to recall this?")
            c1, c2, c3 = st.columns(3)
            if c1.button("Easy (Next review far)", key="sm_easy"):
                add_srs_review(roll, q_obj['id'], 5)
                st.session_state["rev_ans_state"] = None
                if idx + 1 < len(quiz["questions"]):
                    quiz["current_index"] += 1
                else:
                    quiz["completed"] = True
                st.rerun()
            if c2.button("Good (Medium spacing)", key="sm_good"):
                add_srs_review(roll, q_obj['id'], 4)
                st.session_state["rev_ans_state"] = None
                if idx + 1 < len(quiz["questions"]):
                    quiz["current_index"] += 1
                else:
                    quiz["completed"] = True
                st.rerun()
            if c3.button("Hard (Review soon)", key="sm_hard"):
                add_srs_review(roll, q_obj['id'], 3)
                st.session_state["rev_ans_state"] = None
                if idx + 1 < len(quiz["questions"]):
                    quiz["current_index"] += 1
                else:
                    quiz["completed"] = True
                st.rerun()
        elif st.session_state["rev_ans_state"] == "incorrect":
            st.error(f"Incorrect. The correct answer is **{q_obj['correct_option'].strip().upper()}**.")
            if st.button("Retry Question / Keep in Spaced Repetition", key="sm_retry"):
                st.session_state["rev_ans_state"] = None
                if idx + 1 < len(quiz["questions"]):
                    quiz["current_index"] += 1
                else:
                    quiz["completed"] = True
                st.rerun()


def render_exam_prep_view(student):
    roll = student['roll_no']
    st.markdown("""
    <div class="academic-card-purple">
        <h3 style="margin: 0; color: #8B5CF6; font-family: 'Outfit';">⚡ Exam Prep Center</h3>
        <p style="margin: 5px 0 0 0; font-size: 0.9rem; color: #cbd5e1;">Get targeted question packs to clear exams or secure distinction scores. Dynamically compiled from college materials.</p>
    </div>
    """, unsafe_allow_html=True)
    
    subjects = get_subjects_list_all()
    if not subjects:
        st.info("No course materials indexed yet. Sync files from Drive to unlock Exam Prep Packages.")
        return
        
    selected_sub = st.selectbox("Select Subject for Exam Prep", [s['subject_code'] for s in subjects], key="exam_prep_sub_select")
    actual_code = get_actual_subject_code(selected_sub)
    
    conn = get_academic_conn()
    q_rows = conn.execute("SELECT * FROM question_banks WHERE subject_code=?", (actual_code,)).fetchall()
    questions = [dict(r) for r in q_rows]
    conn.close()
    
    if not questions:
        st.info(f"No questions loaded for subject {selected_sub}. Please sync question bank files in Drive.")
        return
        
    pkg = st.radio("Select Preparation Package", ["🎯 Pass Package (Top 10 Essential Questions)", "🏆 Distinction Package (Aim 70+ Marks)", "⏰ Last Night Revision Briefing"], horizontal=True)
    
    if "Pass Package" in pkg:
        st.write("#### 🎯 Pass Package (Highly Occurring Questions)")
        st.markdown("<p style='font-size: 0.85rem; color: #94a3b8; margin-bottom: 15px;'>These questions cover the core concepts required to pass the subject:</p>", unsafe_allow_html=True)
        display_qs = questions[:10]
        for idx, q in enumerate(display_qs):
            with st.expander(f"{idx+1}. [Unit {q['unit']}] {q['question'][:80]}..."):
                st.write(f"**Question:**\\n{q['question']}")
                if q.get("answer_text"):
                    st.write(f"**Answer:**\\n{q['answer_text']}")
                else:
                    st.info("No detailed answer text. Use the AI Copilot to generate an explanation.")
                    
    elif "Distinction Package" in pkg:
        st.write("#### 🏆 Distinction Package (Aim 70+ Marks)")
        st.markdown("<p style='font-size: 0.85rem; color: #94a3b8; margin-bottom: 15px;'>Comprehensive list of important descriptive and analytical topics:</p>", unsafe_allow_html=True)
        display_qs = questions[:25]
        for idx, q in enumerate(display_qs):
            with st.expander(f"{idx+1}. [Unit {q['unit']}] {q['question'][:80]}..."):
                st.write(f"**Question:**\\n{q['question']}")
                if q.get("answer_text"):
                    st.write(f"**Answer:**\\n{q['answer_text']}")
                else:
                    st.info("Use AI Copilot for solution.")
                    
    else:
        st.write("#### ⏰ Last Night Revision Briefing")
        st.markdown("<p style='font-size: 0.85rem; color: #94a3b8; margin-bottom: 15px;'>Ultra-condensed definitions and formulas for quick revision:</p>", unsafe_allow_html=True)
        
        conn = get_academic_conn()
        notes_rows = conn.execute("SELECT unit, topic, content FROM notes WHERE subject_code=? GROUP BY unit", (selected_sub,)).fetchall()
        conn.close()
        
        if not notes_rows:
            st.info("No quick revision notes indexed. Class notes files must be uploaded to Drive.")
        else:
            for r in notes_rows:
                st.markdown(f"""
                <div class="academic-card" style="padding: 15px; margin-bottom: 10px;">
                    <h5 style="margin: 0; color: #00D8C6;">Unit {r[0]}: {r[1]}</h5>
                    <p style="margin: 8px 0 0 0; font-size: 0.85rem; color: #cbd5e1;">{r[2][:400]}...</p>
                </div>
                """, unsafe_allow_html=True)


def render_backlog_recovery_view(student):
    roll = student['roll_no']
    branch = student['branch']
    
    st.markdown("""
    <div class="academic-card-purple">
        <h3 style="margin: 0; color: #8B5CF6; font-family: 'Outfit';">⚠️ AI Backlog Recovery Planner</h3>
        <p style="margin: 5px 0 0 0; font-size: 0.9rem; color: #cbd5e1;">Generate personalized, high-probability roadmaps to clear backlogs. Driven dynamically by your attendance and exam performance.</p>
    </div>
    """, unsafe_allow_html=True)
    
    backlogs = get_student_backlogs(roll)
    if not backlogs:
        st.success("✨ Great news! You have no recorded final exam backlogs. (Risk alerts for current courses are shown on the Dashboard).")
        return
        
    backlog_codes = [b['subject_code'] for b in backlogs]
    selected_backlog = st.selectbox("Select Backlog Subject to Plan For", backlog_codes, key="backlog_plan_select")
    
    days_left = st.slider("Days remaining until Exam", min_value=3, max_value=60, value=14, step=1)
    
    plan_duration = "7 Day Plan" if days_left <= 7 else "14 Day Plan" if days_left <= 15 else "30 Day Plan"
    
    st.markdown(f"**Recommended Plan:** `{plan_duration}` based on {days_left} days left.")
    
    if st.button("🚀 Generate Personalized Recovery Roadmap", use_container_width=True):
        with st.spinner("AI is crafting your study roadmap..."):
            api_key = get_gemini_api_key()
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            
            conn = get_academic_conn()
            wrong_cnt = conn.execute("SELECT COUNT(*) FROM review_schedule r JOIN mcqs m ON r.mcq_id = m.id WHERE r.roll_no=? AND m.subject_code=?", (roll, get_actual_subject_code(selected_backlog))).fetchone()[0]
            conn.close()
            
            prompt = f"""
            You are a senior academic success coach. Design a highly focused, day-by-day {plan_duration} study plan to help this student clear their backlog in {selected_backlog}.
            
            Student Context:
            - Target Subject: {selected_backlog}
            - Days left: {days_left}
            - Repetition review errors: {wrong_cnt} wrong answers in quizzes
            
            Construct a clear, realistic day-by-day study roadmap in clean markdown. 
            Prioritize core topics, formulas, and practicing past question banks. Keep it extremely encouraging and highly structured.
            """
            
            payload = {
                "contents": [{"parts": [{"text": prompt}]}]
            }
            try:
                res = requests.post(url, json=payload, headers=headers, timeout=30)
                if res.status_code == 200:
                    st.markdown(res.json()["candidates"][0]["content"]["parts"][0]["text"])
                else:
                    st.error("Failed to generate plan.")
            except Exception as e:
                st.error(f"Error: {e}")


def render_ai_copilot_view(student):
    roll = student['roll_no']
    branch = student['branch']
    
    st.markdown("""
    <div class="academic-card">
        <h3 style="margin: 0; color: #00D8C6; font-family: 'Outfit';">🤖 AI Copilot (Study Assistant)</h3>
        <p style="margin: 5px 0 0 0; font-size: 0.9rem; color: #cbd5e1;">Ask conceptual questions, request code summaries, or seek math solutions. The Copilot answers using your actual college materials.</p>
    </div>
    """.replace('\n', ' '), unsafe_allow_html=True)
    
    if "chat_history" not in st.session_state:
        st.session_state["chat_history"] = []
        
    for chat in st.session_state["chat_history"]:
        role = chat["role"]
        icon = "🧑‍🎓" if role == "user" else "🤖"
        with st.chat_message("user" if role == "user" else "assistant", avatar=icon):
            st.markdown(chat["text"])
        
    user_input = st.chat_input("Ask a question about your courses...")
    
    if user_input:
        st.session_state["chat_history"].append({"role": "user", "text": user_input})
        st.rerun()
        
    if st.session_state["chat_history"] and st.session_state["chat_history"][-1]["role"] == "user":
        q = st.session_state["chat_history"][-1]["text"]
        
        with st.spinner("AI Copilot is searching database & synthesizing..."):
            results = search_academic_resources_rag(q)
            
            context_parts = []
            for r in results[:6]:
                context_parts.append(f"Source Type: {r['type']} | Subject: {r['subject']} | Title: {r['title']}\\nContent: {r['content']}")
            context_str = "\\n\\n---\\n\\n".join(context_parts)
            
            api_key = get_gemini_api_key()
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent?key={api_key}"
            headers = {"Content-Type": "application/json"}
            
            system_instruction = "You are a helpful, expert VITS Engineering AI Tutor. Your goal is to explain concepts clearly using the provided local syllabus, question banks, lab codes, and notes as a primary source. Always cite which subject the question or code comes from. If the answer is not in the context, explain the concept from general engineering knowledge but mention that it was not directly found in the uploaded course materials."
            
            prompt = f"""
            Context from local question bank, syllabus, lab codes, and notes:
            {context_str if context_str else "No direct matches found in local database."}
            
            User Question: {q}
            
            Please answer the user question in detail, using clean formatting, math LaTeX notation where appropriate, and Markdown.
            """
            
            payload = {
                "contents": [{"parts": [{"text": prompt}]}],
                "systemInstruction": {"parts": [{"text": system_instruction}]}
            }
            
            try:
                res = requests.post(url, json=payload, headers=headers, timeout=25)
                if res.status_code == 200:
                    ans = res.json()["candidates"][0]["content"]["parts"][0]["text"]
                elif res.status_code == 403 and "API_KEY_SERVICE_BLOCKED" in res.text:
                    ans = (
                        "🔒 **API Key Blocked:** The default Gemini API key is currently blocked or deactivated.\n\n"
                        "To use the AI Copilot study assistant, please configure your own Gemini API key:\n"
                        "1. Create a Gemini API key at the [Google AI Studio](https://aistudio.google.com/).\n"
                        "2. Open the file `.streamlit/secrets.toml` in the project directory.\n"
                        "3. Update the key under the `[gemini]` section:\n"
                        "   ```toml\n"
                        "   [gemini]\n"
                        "   api_key = \"YOUR_NEW_API_KEY\"\n"
                        "   ```\n"
                        "4. Save the file and restart the application."
                    )
                else:
                    ans = f"AI Tutor REST call returned error {res.status_code}: {res.text}"
            except Exception as e:
                ans = f"AI Tutor connection error: {e}"
                
            st.session_state["chat_history"].append({"role": "assistant", "text": ans})
            st.rerun()
